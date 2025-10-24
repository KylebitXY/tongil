from django.contrib.auth.models import User
from main1.models import Coach
# tasks.py
# tasks.py
from celery import shared_task
from pdfrw import PdfReader, PdfWriter
import os
from django.conf import settings
import logging

logger = logging.getLogger(__name__)

@shared_task
def fill_pdf_form_task(template_path, output_path, athlete_data):
    try:
        template_pdf = PdfReader(template_path)
        writer = PdfWriter()
        
        for page in template_pdf.pages:
            annotations = page.get('/Annots')
            if annotations:
                for annot in annotations:
                    field = annot.getObject()
                    field_name = field.get('/T')
                    if field_name:
                        field_name = field_name[1:-1]  # Remove parentheses
                        field_value = athlete_data.get(field_name, '')
                        field.update({
                            '/V': field_value
                        })

            writer.addpage(page)
        
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return "PDF form filling successful"
    
    except Exception as e:
        return f"Error: {str(e)}"


from celery import shared_task
from docx import Document
import os
from django.conf import settings
import logging

logger = logging.getLogger(__name__)

@shared_task
def fill_word_form_task(template_path, output_path, athlete_data):
    try:
        # Load the Word document
        doc = Document(template_path)

        # Iterate through paragraphs and replace placeholders
        for para in doc.paragraphs:
            for key, value in athlete_data.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)

        # Iterate through tables and replace placeholders
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in athlete_data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # Save the filled document
        doc.save(output_path)
        
        return "Word form filling successful"
    
    except Exception as e:
        logger.error(f"Error filling Word form: {str(e)}")
        return f"Error: {str(e)}"

def get_coach_for_user(user_id):
    try:
        user = User.objects.get(id=user_id)
        # Find the coach profile associated with this user
        coach = Coach.objects.get(user=user)
        return coach
    except (User.DoesNotExist, Coach.DoesNotExist):
        return None

